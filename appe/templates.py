import os
import uuid
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.pdfbase.pdfmetrics import stringWidth

def generate_html_content(data):
    if not data.get('summary'):
        data['summary'] = f"Adaptive and versatile Frontend Developer with {data.get('experience_years', 0)} years of experience in HTML, CSS, JavaScript, Python"
    
    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Resume</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            margin: 1.2in 1in 1in 1in;
            color: #000000;
            line-height: 1.15;
            font-size: 10.5pt;
        }}
        .container {{
            max-width: 6.27in;
            margin: 0 auto;
        }}
        h1 {{
            font-size: 20pt;
            font-weight: bold;
            margin-bottom: 0.5in;
        }}
        .contact {{
            font-size: 10.5pt;
            margin-bottom: 0.5in;
        }}
        .contact a {{
            color: #0000EE;
            text-decoration: underline;
        }}
        h2 {{
            font-size: 13pt;
            font-weight: bold;
            border-bottom: 1pt solid #000000;
            padding-bottom: 2px;
            margin-bottom: 0.15in;
            margin-top: 0.15in;
            display: inline-block;
            width: auto;
        }}
        .section {{
            margin-bottom: 0.15in;
        }}
        .item {{
            margin-left: 0.25in;
            margin-bottom: 0.15in;
        }}
        ul {{
            list-style-type: disc;
            margin-left: 0.25in;
            padding-left: 0;
            line-height: 1.15;
        }}
        li {{
            margin-bottom: 0.05in;
        }}
        .date {{
            float: right;
            margin-left: 0.5in;
        }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{data.get('name', 'Your Name')}</h1>
        <div class="contact">
            {data.get('location', '')} | {data.get('email', '')} | {data.get('phone', '')} | 
            <a href="{data.get('linkedin', '#')}">{data.get('linkedin', 'LinkedIn Profile')}</a>
        </div>
        <div class="section">
            <h2>Career Objective</h2>
            <p>{data.get('summary', '')}</p>
        </div>
        <div class="section">
            <h2>Educational Background</h2>
"""
    education = ''.join([
        f'<div class="item"><span class="bold">{edu.split(" - ")[0]}</span> <span class="date">{edu.split(" - ")[1] if len(edu.split(" - ")) > 1 else ""}</span><br>' +
        f'<span class="bold">{edu.split(" - ")[2] if len(edu.split(" - ")) > 2 else ""}</span><br>' +
        f'{edu.split(" - ")[3] if len(edu.split(" - ")) > 3 else ""}</div>'
        for edu in data.get('education', '').split('\n') if edu.strip()
    ]) if data.get('education', '').strip() else ''
    skills = ''.join([f'<li>{skill}</li>' for skill in data.get('skills', [])]) if data.get('skills', []) else ''
    experience = ''.join([
        f'<div class="item"><span class="bold">{exp.split(" - ")[0]}</span> - <span class="bold">{exp.split(" - ")[1] if len(exp.split(" - ")) > 1 else ""}</span> <span class="date">{exp.split(" - ")[2] if len(exp.split(" - ")) > 2 else ""}</span><br>' +
        (f'<ul>' + ''.join([f'<li>{ach.strip()}</li>' for ach in exp.split(" - ")[3].split(";") if ach.strip()]) + '</ul>' if len(exp.split(" - ")) > 3 else '') + '</div>'
        for exp in data.get('experience', '').split('\n') if exp.strip()
    ]) if data.get('experience', '').strip() else ''
    projects = ''.join([
        f'<div class="item"><span class="bold">{proj.split(" - ")[0]}</span><br>' +
        f'The language used to build the project is: {proj.split(" - ")[1] if len(proj.split(" - ")) > 1 else ""}<br>' +
        f'Description of the Project: {proj.split(" - ")[2] if len(proj.split(" - ")) > 2 else ""}<br>' +
        f'Link of the Project: <a href="{proj.split(" - ")[3] if len(proj.split(" - ")) > 3 else "#"}">{proj.split(" - ")[3] if len(proj.split(" - ")) > 3 else "No Link"}</a></div>'
        for proj in data.get('projects', '').split('\n\n') if proj.strip()
    ]) if data.get('projects', '').strip() else ''
    certifications = ''.join([
        f'<div class="item"><span class="bold">{cert.split(" - ")[0] if len(cert.split(" - ")) > 0 else cert}</span><br>' +
        (f'<ul>' + ''.join([f'<li>{point.strip()}</li>' for point in cert.split(" - ")[1].split(";") if point.strip()]) + '</ul>' if len(cert.split(" - ")) > 1 else '')
        + '</div>'
        for cert in data.get('certifications', '').split('\n') if cert.strip()
    ]) if data.get('certifications', '').strip() else ''

    html_end = f"""
        </div>
        <div class="section">
            <h2>Skills</h2>
            <ul>{skills}</ul>
        </div>
        <div class="section">
            <h2>Internship Experience</h2>
            {experience}
        </div>
        <div class="section">
            <h2>Projects</h2>
            {projects}
        </div>
        <div class="section">
            <h2>Certifications</h2>
            {certifications}
        </div>
    </div>
</body>
</html>"""

    return html_content + education + html_end

def template(data, output_folder='output', format='pdf'):
    if format == 'pdf':
        os.makedirs(output_folder, exist_ok=True)
        pdf_path = os.path.join(output_folder, f"resume_{uuid.uuid4()}.pdf")
        c = canvas.Canvas(pdf_path, pagesize=A4)
        width = 6.27 * inch  # A4 width minus 2in margins

        # Name
        c.setFont("Helvetica-Bold", 20)
        c.drawString(inch, A4[1] - 1.2*inch, data.get('name', 'Your Name'))
        y = A4[1] - 1.7*inch

        # Contact
        c.setFont("Helvetica", 10.5)
        contact_text = f"{data.get('location', '')} | {data.get('email', '')} | {data.get('phone', '')} | {data.get('linkedin', 'LinkedIn Profile')}"
        c.drawString(inch, y, contact_text)
        y -= 0.5*inch

        # Career Objective
        c.setFont("Helvetica-Bold", 13)
        y -= 0.15*inch
        c.drawString(inch, y, "Career Objective")
        c.line(inch, y - 2, inch + stringWidth("Career Objective", "Helvetica-Bold", 13), y - 2)
        y -= 0.15*inch
        c.setFont("Helvetica", 10.5)
        summary = data.get('summary', '').replace('\n', ' ')
        for line in summary.split('. '):
            if line:
                c.drawString(inch, y, line + '.')
                y -= 0.15*inch
        y -= 0.15*inch

        # Educational Background
        if data.get('education', '').strip():
            c.setFont("Helvetica-Bold", 13)
            y -= 0.15*inch
            c.drawString(inch, y, "Educational Background")
            c.line(inch, y - 2, inch + stringWidth("Educational Background", "Helvetica-Bold", 13), y - 2)
            y -= 0.15*inch
            c.setFont("Helvetica", 10.5)
            for edu in data.get('education', '').split('\n'):
                if edu.strip():
                    parts = edu.split(" - ")
                    if len(parts) > 0:
                        c.setFont("Helvetica-Bold", 10.5)
                        c.drawString(inch, y, parts[0])
                        c.setFont("Helvetica", 10.5)
                        if len(parts) > 1:
                            date_width = stringWidth(parts[1], "Helvetica", 10.5)
                            c.drawString(width - inch - date_width, y, parts[1])
                    if len(parts) > 2:
                        y -= 0.15*inch
                        c.setFont("Helvetica-Bold", 10.5)
                        c.drawString(inch, y, parts[2])
                        c.setFont("Helvetica", 10.5)
                    if len(parts) > 3:
                        y -= 0.15*inch
                        c.drawString(inch, y, parts[3])
                    y -= 0.15*inch
            y -= 0.15*inch

        # Skills
        if data.get('skills', []):
            c.setFont("Helvetica-Bold", 13)
            y -= 0.15*inch
            c.drawString(inch, y, "Skills")
            c.line(inch, y - 2, inch + stringWidth("Skills", "Helvetica-Bold", 13), y - 2)
            y -= 0.15*inch
            c.setFont("Helvetica", 10.5)
            for skill in data.get('skills', []):
                y -= 0.15*inch
                c.drawString(inch + 0.25*inch, y, f"• {skill}")
            y -= 0.15*inch

        # Internship Experience
        if data.get('experience', '').strip():
            c.setFont("Helvetica-Bold", 13)
            y -= 0.15*inch
            c.drawString(inch, y, "Internship Experience")
            c.line(inch, y - 2, inch + stringWidth("Internship Experience", "Helvetica-Bold", 13), y - 2)
            y -= 0.15*inch
            c.setFont("Helvetica", 10.5)
            for exp in data.get('experience', '').split('\n'):
                if exp.strip():
                    parts = exp.split(" - ")
                    if len(parts) > 1:
                        c.setFont("Helvetica-Bold", 10.5)
                        c.drawString(inch, y, f"{parts[0]} - {parts[1]}")
                        c.setFont("Helvetica", 10.5)
                        if len(parts) > 2:
                            date_width = stringWidth(parts[2], "Helvetica", 10.5)
                            c.drawString(width - inch - date_width, y, parts[2])
                    if len(parts) > 3:
                        for ach in parts[3].split(";"):
                            if ach.strip():
                                y -= 0.15*inch
                                c.drawString(inch + 0.25*inch, y, f"• {ach.strip()}")
                    y -= 0.15*inch
            y -= 0.15*inch

        # Projects
        if data.get('projects', '').strip():
            c.setFont("Helvetica-Bold", 13)
            y -= 0.15*inch
            c.drawString(inch, y, "Projects")
            c.line(inch, y - 2, inch + stringWidth("Projects", "Helvetica-Bold", 13), y - 2)
            y -= 0.15*inch
            c.setFont("Helvetica", 10.5)
            for proj in data.get('projects', '').split('\n\n'):
                if proj.strip():
                    parts = proj.split(" - ")
                    y -= 0.15*inch
                    c.setFont("Helvetica-Bold", 10.5)
                    c.drawString(inch, y, parts[0])
                    c.setFont("Helvetica", 10.5)
                    if len(parts) > 1:
                        y -= 0.15*inch
                        c.drawString(inch, y, f"The language used to build the project is: {parts[1]}")
                    if len(parts) > 2:
                        y -= 0.15*inch
                        c.drawString(inch, y, f"Description of the Project: {parts[2]}")
                    if len(parts) > 3:
                        y -= 0.15*inch
                        c.drawString(inch, y, f"Link of the Project: {parts[3]}")
                    y -= 0.25*inch

        # Certifications
        if data.get('certifications', '').strip():
            c.setFont("Helvetica-Bold", 13)
            y -= 0.15*inch
            c.drawString(inch, y, "Certifications")
            c.line(inch, y - 2, inch + stringWidth("Certifications", "Helvetica-Bold", 13), y - 2)
            y -= 0.15*inch
            c.setFont("Helvetica", 10.5)
            for cert in data.get('certifications', '').split('\n'):
                if cert.strip():
                    parts = cert.split(" - ")
                    y -= 0.15*inch
                    c.setFont("Helvetica-Bold", 10.5)
                    c.drawString(inch, y, parts[0])
                    c.setFont("Helvetica", 10.5)
                    if len(parts) > 1:
                        for point in parts[1].split(";"):
                            if point.strip():
                                y -= 0.15*inch
                                c.drawString(inch + 0.25*inch, y, f"• {point.strip()}")
                    y -= 0.25*inch

        # Basic page break check
        if y < 1.2*inch:
            c.showPage()
            y = A4[1] - 1.2*inch

        c.showPage()
        c.save()
        return pdf_path
    elif format == 'docx':
        doc = Document()
        doc.add_heading(data.get('name', 'Your Name'), 0)
        doc.add_paragraph(f"{data.get('location', '')} | {data.get('email', '')} | {data.get('phone', '')} | {data.get('linkedin', 'LinkedIn Profile')}")
        doc.add_heading("Career Objective", level=1)
        doc.add_paragraph(data.get('summary', ''))
        if data.get('education', '').strip():
            doc.add_heading("Educational Background", level=1)
            for edu in data.get('education', '').split('\n'):
                if edu.strip():
                    parts = edu.split(" - ")
                    p = doc.add_paragraph()
                    p.add_run(parts[0]).bold = True
                    p.add_run(" " * 50 + parts[1] if len(parts) > 1 else "").bold = True
                    if len(parts) > 2:
                        doc.add_paragraph(parts[2], style='Heading 3')
                    if len(parts) > 3:
                        doc.add_paragraph(parts[3])
        if data.get('skills', []):
            doc.add_heading("Skills", level=1)
            for skill in data.get('skills', []):
                doc.add_paragraph(f"• {skill}")
        if data.get('experience', '').strip():
            doc.add_heading("Internship Experience", level=1)
            for exp in data.get('experience', '').split('\n'):
                if exp.strip():
                    parts = exp.split(" - ")
                    p = doc.add_paragraph()
                    p.add_run(parts[0]).bold = True
                    p.add_run(" - ").bold = False
                    p.add_run(parts[1]).bold = True
                    p.add_run(" " * 50 + parts[2] if len(parts) > 2 else "").bold = True
                    if len(parts) > 3:
                        for ach in parts[3].split(";"):
                            if ach.strip():
                                doc.add_paragraph(f"• {ach.strip()}")
        if data.get('projects', '').strip():
            doc.add_heading("Projects", level=1)
            for proj in data.get('projects', '').split('\n\n'):
                if proj.strip():
                    parts = proj.split(" - ")
                    doc.add_paragraph(parts[0], style='Heading 3')
                    if len(parts) > 1:
                        doc.add_paragraph(f"The language used to build the project is: {parts[1]}")
                    if len(parts) > 2:
                        doc.add_paragraph(f"Description of the Project: {parts[2]}")
                    if len(parts) > 3:
                        doc.add_paragraph(f"Link of the Project: {parts[3]}")
        if data.get('certifications', '').strip():
            doc.add_heading("Certifications", level=1)
            for cert in data.get('certifications', '').split('\n'):
                if cert.strip():
                    parts = cert.split(" - ")
                    doc.add_paragraph(parts[0], style='Heading 3')
                    if len(parts) > 1:
                        for point in parts[1].split(";"):
                            if point.strip():
                                doc.add_paragraph(f"• {point.strip()}")
        output_path = os.path.join(output_folder, f"resume_{uuid.uuid4()}.docx")
        doc.save(output_path)
        return output_path

def generate_pdf_resume(data, output_folder='output'):
    return template(data, output_folder, format='pdf')

def generate_docx_resume(data, output_folder='output'):
    return template(data, output_folder, format='docx')

# appe/templates.py
def generate_all_previews(data):
    previews = {'template': generate_html_content(data)}  # Return a dictionary with styled HTML
    return previews